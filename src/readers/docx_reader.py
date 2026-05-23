from pathlib import Path

from docx import Document


def extract_docx_text(path: Path) -> tuple[str, dict]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables_text: list[str] = []
    for t_idx, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append(f"--- Tabela {t_idx} ---\n" + "\n".join(rows))

    parts = []
    if paragraphs:
        parts.append("\n\n".join(paragraphs))
    if tables_text:
        parts.append("\n\n".join(tables_text))

    full_text = "\n\n".join(parts)
    metadata = {
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "extraction": "python-docx",
    }
    return full_text, metadata
